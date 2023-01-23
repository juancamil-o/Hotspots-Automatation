import docx
from docx.shared import RGBColor

path = 'C:/Users/jucap/Scripts/Docs/Script.docx' #your docx file path
doc = docx.Document(path)
rosa = RGBColor(255, 0, 255)
naranja = RGBColor(0xff, 0x99, 0x00)
hotspots = []
reviewHotspots = []
cellIndex = 0


def pinkHotspots(index):
    for text in doc.tables[0]._cells[index].paragraphs:
        if not text.runs:
            pass;
        else:
            if text.runs[0].font.color.rgb == rosa:
                hotspots.append(text.text)
    pass
         
    
    
def orangeHotspots(index):
    for text in doc.tables[0]._cells[index].paragraphs:
        if not text.runs:
            pass;
        else:
            if text.runs[0].font.color.rgb == naranja:
                processRun(text)
    pass
            
    
    
def processRun(paragraph):
    foundBold = False
    foundItalic = False
    
    reviewHotspots.append("")
    
    for word in paragraph.runs:
        if  (word.bold == True) and (word.text != " "):
            if ( foundBold == False):
                reviewHotspots[-1] = reviewHotspots[-1] + "<b>"
                foundBold = True      
        elif ( foundBold == True ):
            reviewHotspots[-1] = reviewHotspots[-1] + "</b>"
            foundBold = False
            
        if  (word.italic == True) and (word.text != " "):
            if ( foundItalic == False):
                reviewHotspots[-1] = reviewHotspots[-1] + "<i>"
                foundItalic = True          
        elif ( foundItalic == True ):
            reviewHotspots[-1] = reviewHotspots[-1] + "</i>"
            foundItalic = False
              
            
        reviewHotspots[-1] = reviewHotspots[-1] + str(word.text)
    if(reviewHotspots[-1][0]!='['):
        reviewHotspots.pop()
    return




def organizeHotspots():
    correctedHotspots = []
    correctedHotspots2 = []

    for hp in hotspots:
        correctedHotspots.append(hp[1:])
    for hp in correctedHotspots:
        for index, nhp in reversed(list(enumerate(hp))):
            if nhp == ']':
                correctedHotspots2.append(hp[:index])
                break
    return correctedHotspots2



def organizeHotspots2():
    correctedHotspots = []
    correctedHotspots2 = []

    for hp in reviewHotspots:
        correctedHotspots.append(hp[1:])
    for hp in correctedHotspots:
        for index, nhp in reversed(list(enumerate(hp))):
            if nhp == ']':
                correctedHotspots2.append(hp[:index])
                break
    return correctedHotspots2



def saveFile():
    fileHotspots = organizeHotspots()
    fileReviewHotspots = organizeHotspots2()
    with open('C:/Users/jucap/Scripts/Docs/Hotspots.txt', 'a',encoding="utf-8") as f:
        for line in fileHotspots:
            f.write(f"{line}\n")
    with open('C:/Users/jucap/Scripts/Docs/ReviewHotspots.txt', 'w',encoding="utf-8") as f:
        for line in fileReviewHotspots:
            f.write(f"{line}\n")
      
    
    
def init():
    for index, cells in list(enumerate(doc.tables[0]._cells)):
        if cells.paragraphs[0].text == "AKA":
            orangeHotspots(index)
            pinkHotspots(index)
            saveFile()
        else:
            print(index);    
init()
